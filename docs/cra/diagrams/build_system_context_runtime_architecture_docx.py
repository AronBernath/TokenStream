from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
PNG_PATH = BASE_DIR / "system-context-runtime-architecture.png"
DOCX_PATH = BASE_DIR / "system-context-runtime-architecture.docx"


def font(size, bold=False):
    font_name = "arialbd.ttf" if bold else "arial.ttf"
    candidates = [
        Path("C:/Windows/Fonts") / font_name,
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


FONT_TITLE = font(24, True)
FONT_LABEL = font(19, True)
FONT_TEXT = font(16)
FONT_SMALL = font(14)


def text_size(draw, text, fnt):
    box = draw.textbbox((0, 0), text, font=fnt)
    return box[2] - box[0], box[3] - box[1]


def draw_wrapped_center(draw, rect, text, fnt, fill="#111111", max_chars=22, line_gap=4):
    x1, y1, x2, y2 = rect
    lines = []
    for raw in text.split("\n"):
        lines.extend(wrap(raw, width=max_chars) or [""])
    total_h = sum(text_size(draw, line, fnt)[1] for line in lines) + line_gap * (len(lines) - 1)
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        w, h = text_size(draw, line, fnt)
        draw.text((x1 + ((x2 - x1) - w) / 2, y), line, font=fnt, fill=fill)
        y += h + line_gap


def draw_wrapped_left(draw, pos, text, fnt, fill="#111111", max_chars=46, line_gap=5):
    x, y = pos
    for raw in text.split("\n"):
        for line in wrap(raw, width=max_chars) or [""]:
            draw.text((x, y), line, font=fnt, fill=fill)
            y += text_size(draw, line, fnt)[1] + line_gap
    return y


def box(draw, rect, label, fill="#FFFFFF", outline="#1F2937", width=3, radius=14, max_chars=22):
    draw.rounded_rectangle(rect, radius=radius, fill=fill, outline=outline, width=width)
    draw_wrapped_center(draw, rect, label, FONT_TEXT, max_chars=max_chars)


def datastore(draw, rect, label):
    draw.rounded_rectangle(rect, radius=8, fill="#FAFAFA", outline="#374151", width=3)
    x1, y1, x2, y2 = rect
    draw.line((x1 + 10, y1 + 14, x2 - 10, y1 + 14), fill="#374151", width=2)
    draw.line((x1 + 10, y2 - 14, x2 - 10, y2 - 14), fill="#374151", width=2)
    draw_wrapped_center(draw, rect, label, FONT_SMALL, max_chars=24)


def boundary(draw, rect, title):
    x1, y1, x2, y2 = rect
    dash = 18
    gap = 10
    for x in range(x1, x2, dash + gap):
        draw.line((x, y1, min(x + dash, x2), y1), fill="#2F2F2F", width=4)
        draw.line((x, y2, min(x + dash, x2), y2), fill="#2F2F2F", width=4)
    for y in range(y1, y2, dash + gap):
        draw.line((x1, y, x1, min(y + dash, y2)), fill="#2F2F2F", width=4)
        draw.line((x2, y, x2, min(y + dash, y2)), fill="#2F2F2F", width=4)
    tw, _ = text_size(draw, title, FONT_LABEL)
    draw.rectangle((x1 + 30, y1 - 2, x1 + 50 + tw, y1 + 30), fill="#FFFFFF")
    draw.text((x1 + 40, y1 + 4), title, font=FONT_LABEL, fill="#111111")


def arrow(draw, start, end, label=None, color="#111827", width=3, label_offset=(0, -24)):
    draw.line((start, end), fill=color, width=width)
    sx, sy = start
    ex, ey = end
    dx = ex - sx
    dy = ey - sy
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    head = 16
    left = (ex - ux * head + px * head * 0.55, ey - uy * head + py * head * 0.55)
    right = (ex - ux * head - px * head * 0.55, ey - uy * head - py * head * 0.55)
    draw.polygon([end, left, right], fill=color)
    if label:
        lx = (sx + ex) / 2 + label_offset[0]
        ly = (sy + ey) / 2 + label_offset[1]
        lines = wrap(label, width=26)
        max_w = max(text_size(draw, line, FONT_SMALL)[0] for line in lines)
        total_h = sum(text_size(draw, line, FONT_SMALL)[1] for line in lines) + 4 * (len(lines) - 1)
        draw.rounded_rectangle((lx - 8, ly - 6, lx + max_w + 8, ly + total_h + 6), radius=5, fill="#FFFFFF")
        y = ly
        for line in lines:
            draw.text((lx, y), line, font=FONT_SMALL, fill=color)
            y += text_size(draw, line, FONT_SMALL)[1] + 4


def poly_arrow(draw, points, color="#111827", width=3):
    for start, end in zip(points, points[1:]):
        draw.line((start, end), fill=color, width=width)
    start = points[-2]
    end = points[-1]
    sx, sy = start
    ex, ey = end
    dx = ex - sx
    dy = ey - sy
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    head = 16
    left = (ex - ux * head + px * head * 0.55, ey - uy * head + py * head * 0.55)
    right = (ex - ux * head - px * head * 0.55, ey - uy * head - py * head * 0.55)
    draw.polygon([end, left, right], fill=color)


def render_diagram():
    img = Image.new("RGB", (2400, 1450), "#FFFFFF")
    draw = ImageDraw.Draw(img)

    draw.text((70, 34), "Diagram 1: System Context / Runtime Architecture", font=FONT_TITLE, fill="#111111")
    draw.text((70, 70), "TokenStream CRA risk assessment context view", font=FONT_TEXT, fill="#4B5563")

    boundary(draw, (70, 135, 500, 1310), "External / user-controlled context")
    boundary(draw, (585, 135, 1805, 1310), "TokenStream self-hosted runtime boundary")
    boundary(draw, (1870, 135, 2330, 790), "External providers / tools")

    # External actors and adjacent clients.
    box(draw, (140, 250, 430, 360), "Client application backend\n(machine API consumer)", max_chars=25)
    box(draw, (140, 430, 430, 540), "Human administrator /\noperator", max_chars=22)
    box(draw, (140, 610, 430, 720), "Corpus / source owner", max_chars=22)
    box(draw, (140, 790, 430, 900), "Corpus source systems\nURLs, files, repos,\narchives", max_chars=24)
    box(draw, (140, 970, 430, 1080), "Optional local chat UI\nfor example Open WebUI", max_chars=24)
    box(draw, (140, 1150, 430, 1260), "Client application end user\nno direct access expected", fill="#F9FAFB", max_chars=24)

    # TokenStream runtime services.
    box(draw, (655, 250, 845, 360), "dev-ui\nbrowser admin surface", fill="#F8FAFC", max_chars=20)
    box(draw, (935, 250, 1140, 360), "config_auth\nusers, RBAC, keys,\nregistries", fill="#F8FAFC", max_chars=22)
    box(draw, (1225, 250, 1480, 370), "orchestrator-api\npolicy-aware LLM\nand tool router", fill="#F8FAFC", max_chars=23)
    box(draw, (1545, 250, 1745, 360), "retrieval-api\ncorpus-scoped retrieval", fill="#F8FAFC", max_chars=22)
    box(draw, (1300, 570, 1535, 690), "ingestion-worker\nsource processing\nand indexing", fill="#F8FAFC", max_chars=23)
    box(draw, (1120, 570, 1255, 690), "TEI embedder\nembedding service", fill="#F8FAFC", max_chars=18)

    # Data stores.
    datastore(draw, (655, 870, 900, 980), "config_auth SQLite\nusers, sessions,\nregistry, audit")
    datastore(draw, (950, 870, 1215, 980), "runtime snapshots\nproviders, policies,\nAPI keys, RAG, MCP")
    datastore(draw, (1270, 870, 1515, 980), "object storage / MinIO\nuploaded and\nfetched sources")
    datastore(draw, (1545, 870, 1745, 980), "Qdrant vector store")
    datastore(draw, (1370, 1085, 1745, 1205), "SQLite lexical /\ngraph indexes")

    # External providers.
    box(draw, (1975, 280, 2240, 390), "LLM providers\nOpenAI-compatible,\nDeepSeek, Anthropic,\nlocal", max_chars=23)
    box(draw, (1975, 535, 2240, 645), "MCP servers / tools", max_chars=23)

    # Cross-boundary flows.
    poly_arrow(draw, [(430, 305), (550, 305), (550, 205), (1345, 205), (1345, 250)])
    arrow(draw, (430, 485), (655, 305), None)
    arrow(draw, (430, 665), (655, 305), None)
    arrow(draw, (430, 845), (1300, 635), None)
    arrow(draw, (430, 1025), (1225, 345), None)
    arrow(draw, (270, 1150), (270, 1080), "uses client product UI", label_offset=(20, -5))

    # Internal management/runtime flows.
    arrow(draw, (845, 305), (935, 305), None)
    arrow(draw, (1035, 360), (780, 870), None)
    arrow(draw, (1035, 360), (1085, 870), None)
    arrow(draw, (1225, 355), (1085, 870), None)
    arrow(draw, (1420, 570), (1085, 870), None)
    arrow(draw, (1645, 360), (1085, 870), None)

    # Retrieval and ingestion flows.
    arrow(draw, (1480, 305), (1545, 305), None)
    arrow(draw, (1645, 360), (1645, 870), None)
    arrow(draw, (1645, 360), (1550, 1085), None)
    arrow(draw, (1420, 690), (1390, 870), None)
    arrow(draw, (1300, 635), (1255, 635), None)
    arrow(draw, (1470, 690), (1600, 870), None)
    arrow(draw, (1490, 690), (1550, 1085), None)

    # Provider/tool flows.
    poly_arrow(draw, [(1480, 285), (1830, 285), (1830, 335), (1975, 335)])
    arrow(draw, (1480, 350), (1975, 590), None)

    # Legend.
    legend_x, legend_y = 1880, 900
    draw.text((legend_x, legend_y), "View purpose", font=FONT_LABEL, fill="#111111")
    y = draw_wrapped_left(
        draw,
        (legend_x, legend_y + 44),
        "This is a system context and runtime architecture view. It identifies actors, core components, external systems, major stores, and high-level flows.",
        FONT_SMALL,
        fill="#374151",
        max_chars=48,
    )
    y += 18
    draw.text((legend_x, y), "Not shown here", font=FONT_LABEL, fill="#111111")
    draw_wrapped_left(
        draw,
        (legend_x, y + 44),
        "Individual STRIDE threats, detailed mitigations, protocol-level controls, and the full attack surface register belong in later risk assessment artifacts.",
        FONT_SMALL,
        fill="#374151",
        max_chars=48,
    )

    img.save(PNG_PATH, quality=95)


def set_doc_defaults(doc):
    section = doc.sections[0]
    # Named layout override: diagram artifact uses landscape pages and tighter
    # margins so the embedded architecture view remains readable.
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, before, after in [
        ("Heading 1", 20, 20, 6),
        ("Heading 2", 16, 18, 6),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.font.name = "Arial"
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0, 0, 0)
    r.bold = False

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(subtitle)
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(90, 90, 90)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(11)


def build_docx():
    render_diagram()
    doc = Document()
    set_doc_defaults(doc)
    add_title(
        doc,
        "Diagram 1: System Context / Runtime Architecture",
        "TokenStream CRA risk assessment artifact",
    )

    doc.add_paragraph(
        "This document provides the high-level system context and runtime architecture view used as an input to the TokenStream CRA risk assessment. It identifies the system boundary, expected actors, main runtime components, security-relevant stores, external providers and tools, and the principal high-level flows."
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(PNG_PATH), width=Inches(9.8))

    doc.add_heading("Interpretation", level=1)
    doc.add_paragraph(
        "The diagram should be read as an architecture and context view rather than as the detailed STRIDE model. It establishes what is inside the self-hosted TokenStream runtime boundary and what remains external or user-controlled."
    )
    bullet(doc, "External actors include client application backends, human administrators/operators, corpus/source owners, corpus source systems, optional local chat UI components, and indirect client application end users.")
    bullet(doc, "TokenStream runtime components include dev-ui, config_auth, orchestrator-api, retrieval-api, ingestion-worker, TEI embedder, runtime snapshots, configuration/authentication storage, object storage, Qdrant vectors, and lexical/graph indexes.")
    bullet(doc, "External systems include LLM providers and MCP servers/tools invoked through the orchestrator according to configured policy and available credentials.")

    doc.add_heading("Scope Notes", level=1)
    doc.add_paragraph(
        "This artifact does not enumerate individual STRIDE threats, risk scores, mitigations, vulnerability handling controls, or every attack surface item. Those details should be maintained in the STRIDE Data Flow / Trust Boundary diagram, the Threat Dragon model, and the modelled attack surface register."
    )
    doc.add_paragraph(
        "The optional local chat UI is shown as an external or adjacent client surface because it is not the administrative control plane. Direct user access to dev-ui remains outside the intended access model and should be treated as unintended except for trusted administration."
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
    print(DOCX_PATH)
    print(PNG_PATH)
